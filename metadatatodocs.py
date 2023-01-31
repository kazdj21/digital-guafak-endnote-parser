from docx import Document
from docx.shared import Inches
import pandas as pd
import os

df = pd.read_csv(input("What is the name of the csv file to extract metadata?"))

os.makedirs("outputs")

for item in df.itertuples():
    document = Document()
    document.add_heading(str(item.titles), 0)
    document.add_paragraph(f"Index: {item.Index}")
    document.add_paragraph(f"Contributors: {item.contributors}")
    document.add_paragraph(f"Titles: {item.titles}")
    document.add_paragraph(f"Abstract: {item.abstract}")
    document.add_paragraph(f"Labels: {item.label}")
    document.add_paragraph(f"Periodicals: {item.periodical}")
    document.add_paragraph(f"Keywords: {item.keywords}")
    document.add_paragraph(f"Dates: {item.dates}")
    document.add_paragraph(f"URL: {item.urls}")
    document.save(f"./outputs/{item.Index}.docx")
